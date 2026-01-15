from flask import Flask, render_template, request, redirect, session, send_file, abort
from flask_session import Session
from openpyxl import load_workbook, Workbook
from docx import Document
from datetime import datetime
import os, zipfile, shutil, time, argparse, socket
from threading import Lock

# ---------------- APP SETUP ----------------

app = Flask(__name__)
app.config.from_pyfile("config.py")
app.config["SESSION_TYPE"] = "filesystem"
Session(app)

UPLOAD_FOLDER = "uploads"
EXPORT_FOLDER = "exports"
TEMP_FOLDER = "temp"

os.makedirs(UPLOAD_FOLDER, exist_ok=True)
os.makedirs(EXPORT_FOLDER, exist_ok=True)
os.makedirs(TEMP_FOLDER, exist_ok=True)

quiz_hosted = False
quiz_theme = ""
quiz_time = 0
questions = []

results_lock = Lock()

test_context = {
    "name": "",
    "admin": "",
    "timestamp": "",
    "port": ""
}

# ---------------- UTILITIES ----------------

def get_test_name():
    return test_context.get("name", "")

def get_lan_ip():
    s = socket.socket(socket.AF_INET, socket.SOCK_DGRAM)
    try:
        s.connect(("8.8.8.8", 80))
        ip = s.getsockname()[0]
    except Exception:
        ip = "127.0.0.1"
    finally:
        s.close()
    return ip

def get_latest_file(suffix):
    files = [
        f for f in os.listdir(UPLOAD_FOLDER)
        if f.endswith(suffix) and not f.startswith("~$")
    ]
    return os.path.join(UPLOAD_FOLDER, sorted(files)[-1]) if files else None

def load_questions():
    global questions
    qfile = get_latest_file("_questions.xlsx")
    if not qfile:
        return False

    wb = load_workbook(qfile)
    ws = wb.active
    questions.clear()

    letters = ["A", "B", "C", "D", "E"]
    for row in ws.iter_rows(min_row=2, values_only=True):
        if all(cell == "~" for cell in row):
            break
        options = {
            letters[i]: str(opt)
            for i, opt in enumerate(row[1:6])
            if opt not in (None, "|")
        }
        questions.append({
            "text": row[0],
            "options": options,
            "correct": str(row[6]).strip().upper()
        })
    wb.close()
    return True

def rename_uploaded_files():
    if not all(test_context.values()):
        return

    prefix = f"{test_context['name']}_{test_context['admin']}_{test_context['timestamp']}_{test_context['port']}"

    for base, ext in [("rules", "docx"), ("questions", "xlsx")]:
        old = os.path.join(UPLOAD_FOLDER, f"{base}.{ext}")
        new = os.path.join(UPLOAD_FOLDER, f"{prefix}_{base}.{ext}")
        if os.path.exists(old):
            os.replace(old, new)

# ---------------- ADMIN ----------------

@app.route("/admin", methods=["GET", "POST"])
def admin_login():
    if request.method == "POST":
        if (
            request.form["username"] == app.config["ADMIN_USERNAME"]
            and request.form["password"] == app.config["ADMIN_PASSWORD"]
        ):
            session["admin"] = True
            return redirect("/admin/dashboard")
    return render_template("admin_login.html")

@app.route("/admin/dashboard")
def admin_dashboard():
    if not session.get("admin"):
        abort(403)
    return render_template(
        "admin_dashboard.html",
        lan_ip=get_lan_ip(),
        port=test_context.get("port", "5000")
    )

@app.route("/admin/upload_rules", methods=["POST"])
def upload_rules():
    if not session.get("admin"):
        abort(403)
    request.files["rules"].save(os.path.join(UPLOAD_FOLDER, "rules.docx"))
    return redirect("/admin/dashboard")

@app.route("/admin/upload_questions", methods=["POST"])
def upload_questions():
    global quiz_time
    if not session.get("admin"):
        abort(403)

    path = os.path.join(UPLOAD_FOLDER, "questions.xlsx")
    request.files["questions"].save(path)

    wb = load_workbook(path)
    ws = wb.active
    questions.clear()

    letters = ["A", "B", "C", "D", "E"]
    for row in ws.iter_rows(min_row=2, values_only=True):
        if all(cell == "~" for cell in row):
            break
        options = {
            letters[i]: str(opt)
            for i, opt in enumerate(row[1:6])
            if opt not in (None, "|")
        }
        questions.append({
            "text": row[0],
            "options": options,
            "correct": str(row[6]).strip().upper()
        })
    wb.close()
    quiz_time = len(questions)
    return redirect("/admin/dashboard")

@app.route("/admin/preview", methods=["GET", "POST"])
def preview():
    global quiz_hosted, quiz_theme, quiz_time

    if not session.get("admin"):
        abort(403)

    if request.method == "POST":
        test_context["name"] = request.form["test_name"].replace(" ", "_")
        test_context["admin"] = app.config["ADMIN_USERNAME"]
        test_context["timestamp"] = datetime.now().strftime("%Y%m%d_%H%M%S")
        test_context["port"] = request.environ.get("SERVER_PORT", "5000")

        quiz_time = min(int(request.form["time"]), 180)
        quiz_theme = request.form.get("theme", "")
        quiz_hosted = True

        rename_uploaded_files()
        return redirect("/admin/dashboard")

    rules_lines = []
    if os.path.exists(os.path.join(UPLOAD_FOLDER, "rules.docx")):
        doc = Document(os.path.join(UPLOAD_FOLDER, "rules.docx"))
        rules_lines = [p.text for p in doc.paragraphs]

    return render_template(
        "admin_preview.html",
        questions=questions,
        quiz_time=quiz_time,
        rules_lines=rules_lines,
        theme=quiz_theme
    )

@app.route("/admin/unhost")
def unhost():
    global quiz_hosted, quiz_theme, quiz_time
    if not session.get("admin"):
        abort(403)

    quiz_hosted = False
    quiz_theme = ""
    quiz_time = 0
    questions.clear()
    session.clear()
    return redirect("/admin/dashboard")

# ---------------- USER ----------------

@app.route("/")
def user_entry():
    if not quiz_hosted:
        abort(403)
    return render_template("user_1_details.html", test_name=get_test_name(), theme=quiz_theme)

@app.route("/rules", methods=["GET", "POST"])
def rules():
    if not quiz_hosted:
        abort(403)
    if request.method == "POST":
        session["user"] = dict(request.form)

    rules_file = get_latest_file("_rules.docx")
    if not rules_file:
        abort(500)

    doc = Document(rules_file)
    return render_template(
        "user_2_rules.html",
        rules_lines=[p.text for p in doc.paragraphs],
        theme=quiz_theme,
        test_name=get_test_name()
    )

@app.route("/quiz")
def quiz():
    if not quiz_hosted:
        abort(403)
    if not session.get("user"):
        return redirect("/")
    if not questions:
        load_questions()
    return render_template(
        "user_3_quiz.html",
        questions=questions,
        time=quiz_time * 60,
        theme=quiz_theme,
        user=session["user"],
        test_name=get_test_name()
    )

# ---------------- SUBMIT (FIXED) ----------------

@app.route("/submit", methods=["POST"])
def submit():
    if not quiz_hosted:
        abort(403)

    user = session.get("user")
    if not user:
        abort(400)

    score = 0
    answers = []

    for i, q in enumerate(questions):
        chosen = request.form.get(f"q{i}")
        answers.append(chosen)
        if chosen == q["correct"]:
            score += 1

    prefix = f"{test_context['name']}_{test_context['admin']}_{test_context['timestamp']}_{test_context['port']}"
    wb_path = os.path.join(EXPORT_FOLDER, f"{prefix}_results.xlsx")

    with results_lock:
        if not os.path.exists(wb_path):
            wb = Workbook()
            ws = wb.active
            headers = [
                "Timestamp", "Leader", "Team", "Standard",
                "Phone", "School", "Address",
                "Score", "Time Taken"
            ]
            for idx, q in enumerate(questions, 1):
                headers.append(f"Q{idx}: {q['text']} (Correct: {q['correct']})")
            ws.append(headers)
            wb.save(wb_path)
            wb.close()

        wb = load_workbook(wb_path)
        ws = wb.active
        time_taken = int(request.form.get("time_taken", 0))
        row = [
            datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
            user.get("leader"), user.get("team"),
            user.get("Standard"), user.get("phone"),
            user.get("school"), user.get("address"),
            f"{score}/{len(questions)}",
            f"{time_taken//60:02d}:{time_taken%60:02d}/{quiz_time:02d}:00"
        ]
        row.extend(answers)
        ws.append(row)
        wb.save(wb_path)
        wb.close()

    session.pop("user", None)
    return redirect("/thanks")

@app.route("/thanks")
def thanks():
    if not quiz_hosted:
        abort(403)
    return render_template("user_4_thanks.html", theme=quiz_theme, test_name=get_test_name())

# ---------------- EXPORT ZIP ----------------

@app.route("/admin/export_zip")
def export_zip():
    if not session.get("admin"):
        abort(403)

    prefix = f"{test_context['name']}_{test_context['admin']}_{test_context['timestamp']}_{test_context['port']}"
    zip_path = os.path.join(TEMP_FOLDER, f"{prefix}_EXPORT.zip")

    if os.path.exists(zip_path):
        os.remove(zip_path)

    with zipfile.ZipFile(zip_path, "w", zipfile.ZIP_DEFLATED) as zipf:
        for folder in [UPLOAD_FOLDER, EXPORT_FOLDER]:
            for f in os.listdir(folder):
                if f.startswith(prefix):
                    zipf.write(os.path.join(folder, f), arcname=f)

    return send_file(zip_path, as_attachment=True)

# ---------------- RUN ----------------

if __name__ == "__main__":
    parser = argparse.ArgumentParser()
    parser.add_argument("--host", default="0.0.0.0")
    parser.add_argument("--port", type=int, default=5000)
    args = parser.parse_args()
    app.run(host=args.host, port=args.port, debug=True)
